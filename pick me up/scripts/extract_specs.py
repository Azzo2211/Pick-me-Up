from __future__ import annotations

import json
import sys
from pathlib import Path

import pdfplumber
from docx import Document


def extract_docx(path: Path) -> dict:
    document = Document(path)
    blocks: list[dict] = []

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            blocks.append(
                {
                    "kind": "paragraph",
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else None,
                    "text": text,
                }
            )

    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        blocks.append({"kind": "table", "index": table_index, "rows": rows})

    return {
        "path": str(path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "blocks": blocks,
    }


def extract_pdf(path: Path) -> dict:
    pages = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            pages.append(
                {
                    "page": page_index,
                    "width": page.width,
                    "height": page.height,
                    "text": page.extract_text(x_tolerance=2, y_tolerance=3) or "",
                    "tables": page.extract_tables() or [],
                }
            )
    return {"path": str(path), "page_count": len(pages), "pages": pages}


def main() -> None:
    if len(sys.argv) != 4:
        raise SystemExit("usage: extract_specs.py INPUT.docx INPUT.pdf OUTPUT.json")

    docx_path = Path(sys.argv[1])
    pdf_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = {"docx": extract_docx(docx_path), "pdf": extract_pdf(pdf_path)}
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "docx_paragraphs": payload["docx"]["paragraph_count"],
                "docx_tables": payload["docx"]["table_count"],
                "pdf_pages": payload["pdf"]["page_count"],
                "output": str(output_path),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()

